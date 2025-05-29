# competitor_analysis_reporter.py
import json
import pandas as pd
from datetime import datetime
import statistics
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

class CompetitorReportGenerator:
    def __init__(self, analysis_results):
        """Initialize with analysis results from FacebookCompetitorAnalyzer"""
        self.results = analysis_results
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.output_folder = "output"

    def create_excel_report(self, filename=None):
        """Create comprehensive Excel report with multiple sheets"""
        if not filename:
            filename = f"competitor_analysis_{self.timestamp}.xlsx"
            
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            # Main Overview Sheet
            self._create_overview_sheet(writer)
            
            # Detailed Metrics Sheet
            self._create_detailed_metrics_sheet(writer)
            
            # Engagement Analysis Sheet
            self._create_engagement_sheet(writer)
            
            # Business Analysis Sheet
            self._create_business_sheet(writer)
            
            # Advertising Analysis Sheet
            self._create_advertising_sheet(writer)
            
            # Market Position Sheet
            self._create_market_position_sheet(writer)
            
            # Reel Performance Sheet
            self._create_reel_performance_sheet(writer)
            
        # print(f"Excel report created: {filename}")
        return filename
    
    def _create_overview_sheet(self, writer):
        """Create main overview sheet"""
        competitors = self.results['competitors']
        
        overview_data = []
        for comp in competitors:
            name = comp['page_name']
            metrics = comp['engagement_metrics']
            business = comp['business_analysis']
            ads = comp['advertising_analysis']
            market_pos = self.results['market_position_analysis'].get(name, {})
            
            overview_data.append({
                'Competitor_Name': name,
                'Page_URL': comp['page_url'],
                'Followers': metrics['followers'],
                'Likes': metrics['likes'],
                'Like_to_Follower_Ratio_%': metrics['like_to_follower_ratio'],
                'Engagement_Quality': metrics['follower_engagement_quality'],
                'Avg_Reel_Views': metrics['reel_views']['average_views'],
                'Max_Reel_Views': metrics['reel_views']['max_views'],
                'Content_Performance_Score_%': metrics['content_performance_score'],
                'Total_Reels': metrics['total_reels'],
                'Business_Category': business['categories'],
                'Location': business['location'],
                'Business_Maturity': business['business_maturity'],
                'Contact_Methods_Count': business['contact_diversity_score'],
                'Cross_Platform_Count': business['cross_platform_presence']['total_platforms'],
                'Is_Advertising': ads['is_advertising'],
                'Total_Active_Ads': ads['total_active_ads'],
                'Ad_Intensity': ads['advertising_intensity'],
                'Market_Share_%': market_pos.get('estimated_market_share', 0),
                'Follower_Rank': market_pos.get('follower_rank', 0),
                'Engagement_Rank': market_pos.get('engagement_rank', 0),
                'Competitiveness_Score': market_pos.get('overall_competitiveness', 0)
            })
        
        df_overview = pd.DataFrame(overview_data)
        df_overview.to_excel(writer, sheet_name='Overview', index=False)
        
        # Add summary statistics
        summary_stats = self.results['summary_statistics']
        summary_data = pd.DataFrame([
            ['Total Combined Followers', f"{summary_stats['total_combined_followers']:,}"],
            ['Average Followers', f"{summary_stats['average_followers']:,}"],
            ['Average Content Performance %', f"{summary_stats['average_content_performance']}%"],
            ['Advertising Adoption Rate %', f"{summary_stats['advertising_adoption_rate']}%"],
            ['Cross-Platform Adoption Avg', f"{summary_stats['cross_platform_adoption']:.1f}"],
            ['Analysis Date', self.results['analysis_metadata']['analysis_date'][:10]]
        ], columns=['Metric', 'Value'])
        
        summary_data.to_excel(writer, sheet_name='Summary_Stats', index=False)
    
    def _create_detailed_metrics_sheet(self, writer):
        """Create detailed metrics breakdown"""
        competitors = self.results['competitors']
        
        detailed_data = []
        for comp in competitors:
            name = comp['page_name']
            metrics = comp['engagement_metrics']
            
            detailed_data.append({
                'Competitor': name,
                'Followers': metrics['followers'],
                'Likes': metrics['likes'],
                'Like_Follower_Ratio_%': metrics['like_to_follower_ratio'],
                'Engagement_Quality': metrics['follower_engagement_quality'],
                'Total_Reel_Views': metrics['reel_views']['total_views'],
                'Average_Reel_Views': metrics['reel_views']['average_views'],
                'Median_Reel_Views': metrics['reel_views']['median_views'],
                'Max_Reel_Views': metrics['reel_views']['max_views'],
                'Min_Reel_Views': metrics['reel_views']['min_views'],
                'Total_Reels_Count': metrics['total_reels'],
                'Content_Performance_Score_%': metrics['content_performance_score'],
                'Views_Per_Follower_%': (metrics['reel_views']['average_views'] / metrics['followers'] * 100) if metrics['followers'] > 0 else 0
            })
        
        df_detailed = pd.DataFrame(detailed_data)
        df_detailed.to_excel(writer, sheet_name='Detailed_Metrics', index=False)
    
    def _create_engagement_sheet(self, writer):
        """Create engagement analysis sheet"""
        competitors = self.results['competitors']
        
        engagement_data = []
        for comp in competitors:
            name = comp['page_name']
            metrics = comp['engagement_metrics']
            
            # Calculate engagement ranges
            views = metrics['reel_views']['views_distribution']
            if views:
                high_performing = len([v for v in views if v > metrics['reel_views']['average_views']])
                low_performing = len([v for v in views if v < metrics['reel_views']['average_views']])
            else:
                high_performing = low_performing = 0
            
            engagement_data.append({
                'Competitor': name,
                'Engagement_Quality_Rating': metrics['follower_engagement_quality'],
                'Like_to_Follower_Ratio_%': metrics['like_to_follower_ratio'],
                'Content_Performance_Score': metrics['content_performance_score'],
                'High_Performing_Reels': high_performing,
                'Low_Performing_Reels': low_performing,
                'Consistency_Score': (metrics['reel_views']['median_views'] / metrics['reel_views']['average_views'] * 100) if metrics['reel_views']['average_views'] > 0 else 0,
                'Peak_Performance': metrics['reel_views']['max_views'],
                'Baseline_Performance': metrics['reel_views']['min_views']
            })
        
        df_engagement = pd.DataFrame(engagement_data)
        df_engagement.to_excel(writer, sheet_name='Engagement_Analysis', index=False)
    
    def _create_business_sheet(self, writer):
        """Create business analysis sheet"""
        competitors = self.results['competitors']
        
        business_data = []
        for comp in competitors:
            name = comp['page_name']
            business = comp['business_analysis']
            
            # Parse contact methods
            contact_methods = ', '.join(business['contact_methods']) if business['contact_methods'] else 'None'
            platforms = business['cross_platform_presence']['platforms']
            active_platforms = ', '.join([k for k, v in platforms.items() if v]) if any(platforms.values()) else 'None'
            
            business_data.append({
                'Competitor': name,
                'Business_Category': business['categories'],
                'Location': business['location'],
                'Business_Maturity_Level': business['business_maturity'],
                'Contact_Methods': contact_methods,
                'Contact_Diversity_Score': business['contact_diversity_score'],
                'Cross_Platform_Presence': active_platforms,
                'Total_Platforms': business['cross_platform_presence']['total_platforms'],
                'Integration_Score_%': business['cross_platform_presence']['integration_score'],
                'Has_Physical_Address': 'Yes' if business['location'] != 'not_specified' else 'No',
                'Has_Phone_Contact': 'Yes' if 'phone' in business['contact_methods'] else 'No',
                'Has_WhatsApp_Business': 'Yes' if 'whatsapp' in business['contact_methods'] else 'No'
            })
        
        df_business = pd.DataFrame(business_data)
        df_business.to_excel(writer, sheet_name='Business_Analysis', index=False)
    
    def _create_advertising_sheet(self, writer):
        """Create advertising analysis sheet"""
        competitors = self.results['competitors']
        
        advertising_data = []
        for comp in competitors:
            name = comp['page_name']
            ads = comp['advertising_analysis']
            
            # Parse ad details
            cta_types = ', '.join(ads['cta_types']) if ads['cta_types'] else 'None'
            messaging_themes = ', '.join(ads['ad_messaging_themes']) if ads['ad_messaging_themes'] else 'None'
            
            advertising_data.append({
                'Competitor': name,
                'Currently_Advertising': 'Yes' if ads['is_advertising'] else 'No',
                'Total_Active_Ads': ads['total_active_ads'],
                'Advertising_Intensity': ads['advertising_intensity'],
                'CTA_Types_Used': cta_types,
                'Messaging_Themes': messaging_themes,
                'Ad_Strategy_Diversity': len(ads['ad_messaging_themes']),
                'CTA_Diversity': len(ads['cta_types'])
            })
        
        df_advertising = pd.DataFrame(advertising_data)
        df_advertising.to_excel(writer, sheet_name='Advertising_Analysis', index=False)
    
    def _create_market_position_sheet(self, writer):
        """Create market position analysis sheet"""
        market_data = []
        
        for competitor, position in self.results['market_position_analysis'].items():
            market_data.append({
                'Competitor': competitor,
                'Estimated_Market_Share_%': position['estimated_market_share'],
                'Follower_Rank': position['follower_rank'],
                'Engagement_Rank': position['engagement_rank'],
                'Overall_Competitiveness_Score': position['overall_competitiveness']
            })
        
        df_market = pd.DataFrame(market_data)
        df_market = df_market.sort_values('Overall_Competitiveness_Score', ascending=False)
        df_market.to_excel(writer, sheet_name='Market_Position', index=False)
    
    def _create_reel_performance_sheet(self, writer):
        """Create individual reel performance sheet"""
        reel_data = []
        
        for comp in self.results['competitors']:
            name = comp['page_name']
            views_dist = comp['engagement_metrics']['reel_views']['views_distribution']
            
            for i, views in enumerate(views_dist, 1):
                reel_data.append({
                    'Competitor': name,
                    'Reel_Number': i,
                    'Views': views,
                    'Performance_vs_Average': 'Above' if views > comp['engagement_metrics']['reel_views']['average_views'] else 'Below',
                    'Performance_Score_%': (views / comp['engagement_metrics']['reel_views']['average_views'] * 100) if comp['engagement_metrics']['reel_views']['average_views'] > 0 else 0
                })
        
        df_reels = pd.DataFrame(reel_data)
        df_reels.to_excel(writer, sheet_name='Reel_Performance', index=False)
    
    def create_word_report(self, filename=None):
        """Create formatted Word document report"""
        if not filename:
            filename = f"competitor_analysis_report_{self.timestamp}.docx"
        
        doc = Document()
        
        # Title and Header
        title = doc.add_heading('Facebook Competitor Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Analysis Date: {self.results['analysis_metadata']['analysis_date'][:10]}")
        doc.add_paragraph(f"Total Competitors Analyzed: {self.results['analysis_metadata']['total_competitors']}")
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary_stats = self.results['summary_statistics']
        
        summary_text = f"""
This comprehensive analysis covers {self.results['analysis_metadata']['total_competitors']} Facebook competitors with a combined audience of {summary_stats['total_combined_followers']:,} followers. 

Key findings:
• Average follower count: {summary_stats['average_followers']:,}
• Average content performance: {summary_stats['average_content_performance']}%
• Advertising adoption rate: {summary_stats['advertising_adoption_rate']}%
• Cross-platform integration: {summary_stats['cross_platform_adoption']:.1f} platforms per competitor
        """
        doc.add_paragraph(summary_text)
        
        # Market Leaders
        doc.add_heading('Market Leadership Analysis', level=1)
        insights = self.results['competitive_insights']
        market_leader = insights['market_leader']
        engagement_leader = insights['engagement_leader']
        
        leadership_text = f"""
Market Share Leader: {market_leader[0]} with {market_leader[1]['estimated_market_share']}% estimated market share
Engagement Leader: {engagement_leader[0]} with competitiveness score of {engagement_leader[1]['overall_competitiveness']}
        """
        doc.add_paragraph(leadership_text)
        
        # Individual Competitor Analysis
        doc.add_heading('Individual Competitor Profiles', level=1)
        
        for comp in self.results['competitors']:
            doc.add_heading(comp['page_name'], level=2)
            
            metrics = comp['engagement_metrics']
            business = comp['business_analysis']
            ads = comp['advertising_analysis']
            
            # Create competitor profile table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add data rows
            data_rows = [
                ('Followers', f"{metrics['followers']:,}"),
                ('Likes', f"{metrics['likes']:,}"),
                ('Engagement Quality', metrics['follower_engagement_quality']),
                ('Average Reel Views', f"{metrics['reel_views']['average_views']:,.0f}"),
                ('Content Performance Score', f"{metrics['content_performance_score']}%"),
                ('Business Category', business['categories']),
                ('Location', business['location']),
                ('Business Maturity', business['business_maturity']),
                ('Contact Methods', len(business['contact_methods'])),
                ('Cross-Platform Presence', business['cross_platform_presence']['total_platforms']),
                ('Currently Advertising', 'Yes' if ads['is_advertising'] else 'No'),
                ('Active Ads Count', ads['total_active_ads'])
            ]
            
            for label, value in data_rows:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
        
        # Strategic Opportunities
        doc.add_heading('Strategic Opportunities', level=1)
        
        if insights['advertising_gap']:
            doc.add_heading('Advertising Gaps', level=2)
            gap_text = f"The following {len(insights['advertising_gap'])} competitors are not currently advertising: "
            gap_list = [str(item) for item in insights['advertising_gap'] if item]  # filters out None or empty
            gap_text += ', '.join(gap_list)
            doc.add_paragraph(gap_text)
        
        if insights['content_opportunities']:
            doc.add_heading('Content Performance Gaps', level=2)
            for opp in insights['content_opportunities']:
                doc.add_paragraph(f"• {opp['competitor']}: {opp['performance_gap']}% below market average")
        
        # Market Position Summary Table
        doc.add_heading('Market Position Summary', level=1)
        
        market_table = doc.add_table(rows=1, cols=5)
        market_table.style = 'Table Grid'
        
        # Headers
        header_cells = market_table.rows[0].cells
        headers = ['Competitor', 'Market Share %', 'Follower Rank', 'Engagement Rank', 'Competitiveness Score']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Data rows
        for competitor, position in self.results['market_position_analysis'].items():
            row_cells = market_table.add_row().cells
            row_cells[0].text = competitor if competitor else ""
            row_cells[1].text = f"{position['estimated_market_share']}%"
            row_cells[2].text = str(position['follower_rank'])
            row_cells[3].text = str(position['engagement_rank'])
            row_cells[4].text = str(position['overall_competitiveness'])
        
        doc.save(filename)
        # print(f"Word report created: {filename}")
        return filename
    
    def create_json_report(self, filename=None):
        """Create comprehensive JSON report"""
        if not filename:
            filename = f"competitor_analysis_data_{self.timestamp}.json"
        
        # Add additional calculated fields for better analysis
        enhanced_results = self.results.copy()
        
        # Add performance benchmarks
        all_followers = [comp['engagement_metrics']['followers'] for comp in self.results['competitors']]
        all_views = [comp['engagement_metrics']['reel_views']['average_views'] for comp in self.results['competitors']]
        all_performance = [comp['engagement_metrics']['content_performance_score'] for comp in self.results['competitors']]
        
        benchmarks = {
            'follower_benchmarks': {
                'min': min(all_followers),
                'max': max(all_followers),
                'median': statistics.median(all_followers),
                'q75': statistics.quantiles(all_followers, n=4)[2] if len(all_followers) > 1 else max(all_followers),
                'q25': statistics.quantiles(all_followers, n=4)[0] if len(all_followers) > 1 else min(all_followers)
            },
            'view_benchmarks': {
                'min': min(all_views),
                'max': max(all_views),
                'median': statistics.median(all_views),
                'q75': statistics.quantiles(all_views, n=4)[2] if len(all_views) > 1 else max(all_views),
                'q25': statistics.quantiles(all_views, n=4)[0] if len(all_views) > 1 else min(all_views)
            },
            'performance_benchmarks': {
                'min': min(all_performance),
                'max': max(all_performance),
                'median': statistics.median(all_performance),
                'q75': statistics.quantiles(all_performance, n=4)[2] if len(all_performance) > 1 else max(all_performance),
                'q25': statistics.quantiles(all_performance, n=4)[0] if len(all_performance) > 1 else min(all_performance)
            }
        }
        
        enhanced_results['benchmarks'] = benchmarks
        
        # Add quick insights
        quick_insights = {
            'top_performer_by_followers': max(self.results['competitors'], key=lambda x: x['engagement_metrics']['followers'])['page_name'],
            'top_performer_by_engagement': max(self.results['competitors'], key=lambda x: x['engagement_metrics']['content_performance_score'])['page_name'],
            'most_active_advertiser': max(self.results['competitors'], key=lambda x: x['advertising_analysis']['total_active_ads'])['page_name'],
            'most_cross_platform': max(self.results['competitors'], key=lambda x: x['business_analysis']['cross_platform_presence']['total_platforms'])['page_name'],
            'newest_competitor': min(self.results['competitors'], key=lambda x: int(x['extraction_data'].get('Creation_date', '0')))['page_name'] if self.results['competitors'] else 'N/A'
        }
        
        enhanced_results['quick_insights'] = quick_insights
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(enhanced_results, f, ensure_ascii=False, indent=2)
        
        # print(f"JSON report created: {filename}")
        return filename
    
    def generate_all_reports(self, base_filename=None):
        """Generate all three report formats"""
        if not base_filename:
            base_filename = f"competitor_analysis_{self.timestamp}"
        
        files_created = {}
        
        # Create all reports
        files_created['excel'] = self.create_excel_report(f"{base_filename}.xlsx")
        files_created['word'] = self.create_word_report(f"{base_filename}.docx")
        files_created['json'] = self.create_json_report(f"{base_filename}.json")
        
        print("\n" + "="*60)
        print("📊 REPORT GENERATION COMPLETE")
        print("="*60)
        print(f"Excel Report: {files_created['excel']}")
        print(f"Word Report: {files_created['word']}")
        print(f"JSON Report: {files_created['json']}")
        print("\nAll reports contain comprehensive analysis data that can be")
        
        return files_created
